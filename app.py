import sqlite3
from docx import Document
from docx.shared import Pt
from flask import Flask, render_template, request

app = Flask(__name__)


def get_db_connection():
    conn = sqlite3.connect('database.db')
    conn.row_factory = sqlite3.Row
    return conn


@app.route('/')
def index():
    conn = get_db_connection()
    tables = conn.execute("SELECT name FROM sqlite_master WHERE type='table';")
    print(conn.execute('SELECT * FROM Собеседование2').fetchall())
    return render_template('index.html', tables=tables)


@app.after_request
def add_header(r):
    """
    Add headers to both force latest IE rendering engine or Chrome Frame,
    and also to cache the rendered page for 10 minutes.
    """
    r.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    r.headers["Pragma"] = "no-cache"
    r.headers["Expires"] = "0"
    r.headers['Cache-Control'] = 'public, max-age=0'
    return r


def createDocx(columns, workflow):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    document.add_paragraph(workflow)
    for i in range(1, len(columns)):
        document.add_paragraph(columns[i] + ': {{' + columns[i] + '}}')
    document.save(workflow + '.docx')


def fillDocx(data, workflow):
    print(data)
    template = workflow + '.docx'
    document = Document(template)
    for key, value in data:
        for paragraph in document.paragraphs:
            replace_text(paragraph, key, value)
    document.save(workflow + '.docx')


def replace_text(paragraph, key, value):
    if key in paragraph.text:
        key = '{{' + key + '}}'
        paragraph.text = paragraph.text.replace(key, value)


@app.route('/runworkflow/<string:workflow>', methods=["POST", "GET"])
def runworkflow(workflow):
    if request.method == 'POST':
        data = request.form.to_dict()
        values = []
        for key, value in data.items():
            values.append(value)
        print(values)
        conn = get_db_connection()
        sql_scheme = 'INSERT INTO ' + workflow + '('
        counter = 0
        for key in data:
            counter += 1
            if counter == len(data):
                sql_scheme += key + ')'
            else:
                sql_scheme += key + ','
        sql_scheme += ' VALUES('
        for i in range(0, counter):
            if i == counter - 1:
                sql_scheme += '?)'
            else:
                sql_scheme += '?,'
        conn.execute(sql_scheme, tuple(values))
        print(sql_scheme)
        print(values)
        conn.commit()
        conn.close()
        fillDocx(data.items(), workflow)
    conn = get_db_connection()
    variables = conn.execute('SELECT * FROM ' + workflow)
    return render_template('runworkflow.html', variables=variables, workflow=workflow)


@app.route('/<string:workflowName>', methods=["POST", "GET"])
def workflow(workflowName):
    print(workflowName)
    conn = get_db_connection()
    column_names = conn.execute('SELECT * FROM ' + workflowName)
    return render_template('workflow.html', workflowName=workflowName, columns=column_names)


@app.route('/createTemplate', methods=["POST", "GET"])
def createTemplate():
    if request.method == 'POST':
        variables = request.form.getlist('skill[]')
        print(variables)
        conn = get_db_connection()
        sql_scheme = 'CREATE TABLE ' + variables[0].replace(" ", "") + '(\n'
        for i in range(1, len(variables)):
            if (i == len(variables) - 1):
                sql_scheme = sql_scheme + variables[i].replace(" ", "") + ' VARCHAR(255) NOT NULL\n'
            else:
                sql_scheme = sql_scheme + variables[i].replace(" ", "") + ' VARCHAR(255) NOT NULL,\n'
        sql_scheme = sql_scheme + ');'
        print(sql_scheme)
        conn.execute(sql_scheme)
        conn.close()
        createDocx(variables, variables[0])

    return render_template('createWorkFlow.html')


if __name__ == '__main__':
    app.run()
